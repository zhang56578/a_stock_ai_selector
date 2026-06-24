#!/usr/bin/env python3
"""
A股智选助手 v2.0 - AI驱动选股与持仓优化平台
完整功能仪表板 | 解决股民五大痛点
运行: streamlit run app.py
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime, timedelta, time as dt_time
import io
import json
import hashlib
import time as _time

try:
    from streamlit_autorefresh import st_autorefresh
    HAS_AUTOREFRESH = True
except ImportError:
    HAS_AUTOREFRESH = False

from data_provider import (
    get_kline, get_fund_flow, get_sector_hotspots,
    get_us_stock_linkage, get_macro_prediction,
    add_technical_indicators, get_market_overview,
    get_limit_up_data, get_northbound_flow,
    get_financial_summary, get_news,
)
from ai_engine import AIEngine
from technical_indicators import TechnicalIndicators
from backtest_engine import (
    BacktestEngine, BUILTIN_STRATEGIES,
    run_backtest, compare_strategies,
)
from stock_scanner import (
    STOCK_POOLS, scan_pool, scan_all_pools, scan_stocks,
    filter_buy_signals, summarize_by_sector,
)
from watchlist import (
    load_watchlist, add_to_watchlist, remove_from_watchlist,
    is_in_watchlist, get_watchlist_pool, get_watchlist_count,
    get_watchlist_df, clear_watchlist,
)

# ==================== 页面配置 ====================
st.set_page_config(
    page_title="A股智选 - AI选股工具",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ==================== CSS 样式 ====================
st.markdown("""
<style>
    .main-header { font-size: 2em; font-weight: 700; color: #1f77b4; margin-bottom: 0.5em; }
    .metric-card {
        background: #f8f9fa; border-radius: 10px; padding: 16px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center;
    }
    .signal-buy { color: #ff4d4d; font-weight: bold; }
    .signal-sell { color: #00cc00; font-weight: bold; }
    .risk-high { color: #ff4444; }
    .risk-mid { color: #ffaa00; }
    .risk-low { color: #00cc00; }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] { border-radius: 4px 4px 0px 0px; padding: 10px 16px; }
</style>
""", unsafe_allow_html=True)


# ==================== 初始化缓存 ====================
@st.cache_data(ttl=300, show_spinner=False)
def cached_get_kline(code, days):
    return get_kline(code, days)


@st.cache_data(ttl=120, show_spinner=False)
def cached_get_sector_hotspots():
    return get_sector_hotspots()


@st.cache_data(ttl=180, show_spinner=False)
def cached_get_us_linkage():
    return get_us_stock_linkage()


@st.cache_data(ttl=60, show_spinner=False)
def cached_get_fund_flow(code):
    return get_fund_flow(code)


# ==================== 股票选择配置 ====================
STOCK_POOL = {
    "688017": {"name": "绿的谐波/半导体设备", "sector": "半导体设备"},
    "688981": {"name": "中芯国际/存储芯片", "sector": "存储芯片"},
    "688012": {"name": "中微公司/半导体设备", "sector": "半导体设备"},
    "002049": {"name": "紫光国微/存储芯片", "sector": "存储芯片"},
    "002371": {"name": "北方华创/半导体设备", "sector": "半导体设备"},
    "300474": {"name": "景嘉微/GPU", "sector": "AI算力"},
    "688256": {"name": "寒武纪/AI芯片", "sector": "AI算力"},
    "300502": {"name": "新易盛/光模块", "sector": "AI算力"},
    "300750": {"name": "宁德时代/新能源", "sector": "新能源"},
    "601012": {"name": "隆基绿能/光伏", "sector": "新能源"},
    "000001": {"name": "平安银行", "sector": "银行"},
    "600519": {"name": "贵州茅台", "sector": "白酒"},
    "300059": {"name": "东方财富", "sector": "券商"},
    "002415": {"name": "海康威视/AI", "sector": "AI"},
    "002230": {"name": "科大讯飞/AI", "sector": "AI"},
    "600641": {"name": "万业企业/离子注入", "sector": "长鑫存储概念"},
    "688361": {"name": "中科飞测/检测", "sector": "长鑫存储概念"},
    "688126": {"name": "沪硅产业/硅片", "sector": "长鑫存储概念"},
    "688147": {"name": "微导纳米/ALD", "sector": "长鑫存储概念"},
    "688037": {"name": "芯源微/Track", "sector": "长鑫存储概念"},
    "000636": {"name": "风华高科/MLCC龙头", "sector": "MLCC概念"},
    "300408": {"name": "三环集团/MLCC", "sector": "MLCC概念"},
    "603678": {"name": "火炬电子/特容", "sector": "MLCC概念"},
    "002156": {"name": "通富微电/先进封测", "sector": "先进封装(Chiplet)"},
    "603005": {"name": "晶方科技/WLCSP", "sector": "先进封装(Chiplet)"},
    "600460": {"name": "士兰微/GaN", "sector": "氮化镓GaN"},
    "600703": {"name": "三安光电/GaN+SiC", "sector": "氮化镓GaN"},
    "603290": {"name": "斯达半导/SiC模组", "sector": "碳化硅SiC"},
    "002023": {"name": "海特高新/InP", "sector": "磷化铟InP"},
    "603773": {"name": "沃格光电/玻璃基板", "sector": "玻璃基板"},
    "301071": {"name": "力量钻石/金刚石", "sector": "人造钻石"},
}

# 信号中英文映射
SIGNAL_CN = {
    'STRONG_BUY': '🟢强烈买入', 'BUY': '🔵买入', 'WEAK_BUY': '⚪偏多',
    'NEUTRAL': '⬜中性', 'WEAK_SELL': '🟡偏空', 'SELL': '🟠卖出',
    'STRONG_SELL': '🔴强烈卖出',
}
# 列名中英文映射（扫描结果表）
COL_CN = {
    'code': '代码', 'name': '名称', 'sector': '板块', 'price': '现价',
    'change%': '当日涨跌', 'signal': '信号', 'buy_score': '买入分',
    'sell_score': '卖出分', 'rsi': 'RSI',
    'buy_strategy': '🟢买入策略', 'sell_strategy': '🔴卖出策略',
    'buy_reasons': '买入原因', 'sell_reasons': '卖出原因',
    'ret_20d': '20日涨跌', 'ret_60d': '60日涨跌',
    'macd_cross_wr': 'MACD金叉胜率',
}

def _label_reasons(reasons_str: str) -> str:
    """给买入原因加策略标签（与回测5策略一一对应）"""
    if not reasons_str:
        return ''
    parts = reasons_str.split(';')
    labeled = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # MACD金叉+放量
        if 'MACD' in p and ('金叉' in p or '底背离' in p or '红柱' in p):
            labeled.append(f'[MACD金叉+放量] {p}')
        elif '放量' in p and ('上涨' in p or '确认' in p):
            labeled.append(f'[MACD金叉+放量] {p}')
        # 均线系统
        elif '均线' in p and ('多头' in p or '排列' in p):
            labeled.append(f'[均线系统] {p}')
        elif 'MA' in p and ('上穿' in p or '偏离' in p):
            labeled.append(f'[均线系统] {p}')
        elif '偏离MA' in p:
            labeled.append(f'[均线系统] {p}')
        # 布林带反转
        elif '布林' in p or 'BB' in p:
            labeled.append(f'[布林带反转] {p}')
        elif 'RSI' in p and ('超卖' in p or '≤' in p):
            labeled.append(f'[布林带反转] {p}')
        elif 'KDJ' in p and ('J值' in p or '超卖' in p):
            labeled.append(f'[布林带反转] {p}')
        elif 'WR' in p and '超卖' in p:
            labeled.append(f'[布林带反转] {p}')
        elif 'CCI' in p:
            labeled.append(f'[布林带反转] {p}')
        # 突破放量
        elif '涨幅' in p and '>' in p:
            labeled.append(f'[突破放量] {p}')
        # 多因子综合 — 兜底
        else:
            labeled.append(f'[多因子综合] {p}')
    return '; '.join(labeled)


# ==================== K线可视化 ====================
def plot_kline(df: pd.DataFrame, title: str = ""):
    fig = make_subplots(
        rows=4, cols=1, shared_xaxes=True,
        vertical_spacing=0.025,
        row_heights=[0.48, 0.17, 0.17, 0.18],
        subplot_titles=(title, 'MACD', 'KDJ / RSI', '成交量'),
    )

    # K线 + 均线
    fig.add_trace(go.Candlestick(
        x=df.index, open=df['open'], high=df['high'],
        low=df['low'], close=df['close'], name='K线',
        increasing_line_color='#ff4d4d', decreasing_line_color='#00cc00',
    ), row=1, col=1)

    for col in ['MA5', 'MA10', 'MA20', 'MA60']:
        if col in df.columns:
            fig.add_trace(go.Scatter(
                x=df.index, y=df[col], name=col,
                line=dict(width=1.2),
            ), row=1, col=1)

    # 布林带
    if 'BB_UPPER' in df.columns:
        fig.add_trace(go.Scatter(
            x=df.index, y=df['BB_UPPER'], name='BB上轨',
            line=dict(width=0.5, color='gray', dash='dash'),
        ), row=1, col=1)
        fig.add_trace(go.Scatter(
            x=df.index, y=df['BB_LOWER'], name='BB下轨',
            line=dict(width=0.5, color='gray', dash='dash'),
            fill='tonexty', fillcolor='rgba(128,128,128,0.05)',
        ), row=1, col=1)

    # MACD
    fig.add_trace(go.Scatter(
        x=df.index, y=df['MACD'], name='MACD', line=dict(color='blue'),
    ), row=2, col=1)
    fig.add_trace(go.Scatter(
        x=df.index, y=df['Signal'], name='Signal', line=dict(color='orange'),
    ), row=2, col=1)
    fig.add_trace(go.Bar(
        x=df.index, y=df['Histogram'], name='柱',
        marker_color=df['Histogram'].apply(lambda x: '#ff4d4d' if x >= 0 else '#00cc00'),
    ), row=2, col=1)

    # KDJ + RSI
    if all(k in df.columns for k in ['K', 'D', 'J']):
        fig.add_trace(go.Scatter(
            x=df.index, y=df['K'], name='K', line=dict(color='#ff4d4d', width=1),
        ), row=3, col=1)
        fig.add_trace(go.Scatter(
            x=df.index, y=df['D'], name='D', line=dict(color='#1f77b4', width=1),
        ), row=3, col=1)
        fig.add_trace(go.Scatter(
            x=df.index, y=df['J'], name='J', line=dict(color='#9467bd', width=0.8),
        ), row=3, col=1)
    fig.add_hline(y=80, line_dash="dash", line_color="red", opacity=0.4, row=3, col=1)
    fig.add_hline(y=20, line_dash="dash", line_color="green", opacity=0.4, row=3, col=1)

    # RSI on same subplot as secondary axis
    if 'RSI' in df.columns:
        fig.add_trace(go.Scatter(
            x=df.index, y=df['RSI'], name='RSI',
            line=dict(color='purple', dash='dot', width=1),
        ), row=3, col=1)

    # 成交量
    colors = ['#ff4d4d' if df['close'].iloc[i] >= df['open'].iloc[i] else '#00cc00' for i in range(len(df))]
    fig.add_trace(go.Bar(
        x=df.index, y=df['volume'], name='成交量', marker_color=colors, opacity=0.7,
    ), row=4, col=1)
    if 'VOLUME_MA20' in df.columns:
        fig.add_trace(go.Scatter(
            x=df.index, y=df['VOLUME_MA20'], name='量MA20', line=dict(color='orange', width=1),
        ), row=4, col=1)

    fig.update_layout(
        height=800, showlegend=True, hovermode='x unified',
        xaxis_rangeslider_visible=False,
        margin=dict(l=0, r=0, t=40, b=0),
    )
    fig.update_xaxes(rangeslider_visible=False)
    return fig


def plot_chip(df: pd.DataFrame):
    """筹码分布图"""
    ti = TechnicalIndicators(df.copy())
    chip = ti.get_chip_distribution()
    if not chip:
        st.warning("数据不足，无法绘制筹码分布")
        return

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=chip['prices'], y=chip['chips'], name='筹码',
        marker_color='rgba(30, 144, 255, 0.6)', orientation='h',
    ))
    fig.add_hline(
        y=chip['current_price'], line_dash="dash",
        line_color="red", annotation_text=f"现价 {chip['current_price']:.2f}",
    )
    fig.add_hline(
        y=chip['avg_cost'], line_dash="dot",
        line_color="orange", annotation_text=f"均成本 {chip['avg_cost']:.2f}",
    )
    fig.update_layout(
        title="筹码分布图", height=400,
        yaxis_title="价格", xaxis_title="筹码数量",
        margin=dict(l=0, r=0, t=40, b=0),
    )
    st.plotly_chart(fig, use_container_width=True)


def plot_signal_gauge(buy_score, sell_score):
    """买卖信号仪表图"""
    fig = go.Figure()
    net = buy_score - sell_score
    color = '#ff4d4d' if net > 0 else '#00cc00'

    fig.add_trace(go.Indicator(
        mode="gauge+number+delta",
        value=net,
        delta={'reference': 0},
        title={'text': "综合信号强度"},
        gauge={
            'axis': {'range': [-100, 100]},
            'bar': {'color': color},
            'steps': [
                {'range': [-100, -30], 'color': 'rgba(0,200,0,0.15)'},
                {'range': [-30, 30], 'color': 'rgba(200,200,200,0.15)'},
                {'range': [30, 100], 'color': 'rgba(255,0,0,0.15)'},
            ],
            'threshold': {
                'line': {'color': "black", 'width': 2},
                'thickness': 0.8, 'value': net,
            },
        },
    ))
    fig.update_layout(height=250, margin=dict(l=10, r=10, t=40, b=10))
    st.plotly_chart(fig, use_container_width=True)


# ==================== DOCX报告导出 ====================
def create_docx_report(report_text: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('| ') and ' | ' in line:
            doc.add_paragraph(line if '---' not in line else '')
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        elif line.startswith('---'):
            pass
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def is_market_open() -> bool:
    """判断当前是否在A股交易时间"""
    now = datetime.now()
    if now.weekday() >= 5:
        return False
    t = now.time()
    return (dt_time(9, 30) <= t <= dt_time(11, 30)) or (dt_time(13, 0) <= t <= dt_time(15, 0))


def fmt_market_status() -> str:
    if is_market_open():
        return "🟢 交易中"
    now = datetime.now()
    if now.weekday() >= 5:
        return "⚫ 休市(周末)"
    t = now.time()
    if t < dt_time(9, 30):
        return "🟡 盘前"
    if dt_time(15, 0) < t:
        return "🔴 已收盘"
    return "🟡 午休"


def getProgressPct(current, total, cap=True):
    """计算百分比，可选封顶100"""
    if not total or total <= 0:
        return 0
    pct = min(100, round(current / total * 100)) if cap else round(current / total * 100)
    return pct


# ==================== 主界面 ====================
def main():
    # 头部
    col_title, col_status = st.columns([4, 1])
    with col_title:
        st.markdown('<div class="main-header">📈 A股智选助手 v2.0 - AI驱动选股 & 持仓优化</div>',
                    unsafe_allow_html=True)
    with col_status:
        st.caption(f"🕐 {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    st.caption("覆盖技术验证 | 持仓优化 | 价值研判 | 行情复盘 | 策略回测 — 解决股民五大核心痛点")

    # ==================== 侧边栏 ====================
    with st.sidebar:
        st.header("⚙️ 控制面板")

        st.subheader("📌 股票选择")
        selected_label = st.selectbox(
            "快速选择",
            options=["自定义输入"] + [f"{k} {v['name']}" for k, v in STOCK_POOL.items()],
            index=0,
        )
        if selected_label == "自定义输入":
            custom = st.text_input("输入股票代码", "688017", key="custom_code")
            selected_code = custom.zfill(6)
            st.info("示例：688981(存储芯片) | 000001(平安银行) | 300474(GPU)")
        else:
            selected_code = selected_label.split()[0]

        stock_info = STOCK_POOL.get(selected_code, {"name": "未知", "sector": "未分类"})

        st.subheader("📊 分析参数")
        analysis_days = st.slider("K线周期(日)", 30, 365, 120)
        use_real_data = st.checkbox("使用真实数据(需网络)", value=True, help="取消则使用模拟数据")
        show_all_ta = st.checkbox("显示全部技术指标", value=False, help="勾选后计算全部指标(较慢)")

        st.divider()

        st.subheader("💼 我的持仓")
        with st.expander("持仓配置"):
            holdings_df = pd.DataFrame([
                {"code": "688017", "name": "绿的谐波", "cost": 120.5, "shares": 500, "sector": "半导体设备"},
                {"code": "688981", "name": "存储芯片", "cost": 42.3, "shares": 800, "sector": "存储芯片"},
            ])
            edited = st.data_editor(
                holdings_df, num_rows="dynamic",
                column_config={
                    "code": "代码", "name": "名称", "cost": "成本价(元)",
                    "shares": "股数", "sector": "板块",
                },
            )
            total_capital_input = st.number_input("总资金(万元)", 1, 10000, 50, step=1)

        st.divider()

        st.subheader("🤖 AI设置")
        use_llm = st.checkbox("启用LLM增强", value=False, help="需要配置API Key")
        if use_llm:
            llm_provider = st.selectbox("LLM提供商", ["openai", "gemini"])
            api_key = st.text_input("API Key", type="password")
        else:
            llm_provider = "openai"
            api_key = None

        if st.button("🔄 重新分析", use_container_width=True):
            st.cache_data.clear()
            st.rerun()

    # ==================== 🐉 侧边栏：一键擒龙快速入口 ====================
    with st.sidebar:
        st.divider()
        st.markdown("### 🐉 一键擒龙")
        st.caption("全市场AI扫描，秒出买入信号")

        dragon_pool = st.selectbox(
            "扫描范围",
            options=["全部板块"] + ["⭐ 我的关注"] + list(STOCK_POOLS.keys()),
            index=0,
            key="dragon_pool_sidebar",
        )
        dragon_days = st.slider("分析天数", 30, 200, 90, 30, key="dragon_days_sidebar")
        dragon_workers = st.slider("并发线程", 1, 8, 5, 1, key="dragon_workers_sidebar")
        exclude_star = st.checkbox("排除科创板(688开头)", value=True, key="dragon_exclude_star",
                                   help="科创板需要50万门槛，无权限者建议勾选")

        if st.button("🚀 开始擒龙扫描", type="primary", use_container_width=True, key="dragon_scan_sidebar"):
            with st.spinner("⏳ 扫描中..."):
                all_stocks = []
                if dragon_pool == "全部板块":
                    for sector, stocks in STOCK_POOLS.items():
                        for code, name in stocks:
                            if not (exclude_star and code.startswith('688')):
                                all_stocks.append((code, name, sector))
                    for s in load_watchlist():
                        all_stocks.append((s['code'], s.get('name',''), s.get('sector','用户关注')))
                elif dragon_pool == "⭐ 我的关注":
                    for s in load_watchlist():
                        all_stocks.append((s['code'], s.get('name',''), s.get('sector','用户关注')))
                else:
                    for code, name in STOCK_POOLS.get(dragon_pool, []):
                        if not (exclude_star and code.startswith('688')):
                            all_stocks.append((code, name, dragon_pool))

                scan_df = scan_stocks(all_stocks, days=dragon_days, max_workers=dragon_workers)
                if not scan_df.empty:
                    buy_df = filter_buy_signals(scan_df, min_buy_score=10)
                    st.session_state['scan_df'] = scan_df
                    st.session_state['buy_df'] = buy_df
                    st.session_state['scan_time'] = datetime.now().strftime('%H:%M:%S')
                    st.success(f"✅ 扫描完成: {len(buy_df)}只买入 / {len(scan_df)}只全部")
                    st.caption(f"最强: {scan_df.groupby('sector')['buy_score'].mean().idxmax() if not scan_df.empty else '-'}")
                    st.caption("👉 到「🐉 一键擒龙」Tab查看详情")
                else:
                    st.warning("扫描无结果")

        # 关注列表入口
        wl_count = get_watchlist_count()
        st.caption(f"📋 已关注: {wl_count} 只")
        with st.expander("📋 管理关注列表"):
            if wl_count == 0:
                st.caption("暂无关注股票。扫描后点击「⭐ 加入关注」添加")
            else:
                wl_df = get_watchlist_df()
                for _, row in wl_df.iterrows():
                    c1, c2 = st.columns([4, 1])
                    with c1:
                        st.write(f"{row['code']} {row.get('name','')} [{row.get('sector','')}]")
                    with c2:
                        if st.button("🗑️", key=f"wl_del_{row['code']}"):
                            remove_from_watchlist(row['code'])
                            st.rerun()
                if st.button("🗑️ 清空全部关注"):
                    clear_watchlist()
                    st.rerun()

    # ==================== 🤖 侧边栏：实时刷新 ====================
    with st.sidebar:
        st.divider()
        st.markdown(f"### ⏱ 实时状态 {fmt_market_status()}")

        if HAS_AUTOREFRESH:
            auto_refresh = st.checkbox("🔄 自动刷新", value=is_market_open(),
                                       key="auto_refresh", help="勾选后自动定时刷新数据")
            if auto_refresh:
                refresh_interval = st.slider("间隔(秒)", 10, 120, 30, 10, key="refresh_interval")
                st_autorefresh(interval=refresh_interval * 1000, key="global_ar", limit=None)
        else:
            auto_refresh = st.checkbox("🔄 自动刷新(需安装streamlit-autorefresh)", value=False,
                                       key="auto_refresh", disabled=True)
            st.caption("`pip install streamlit-autorefresh`")

        if st.button("🔃 手动刷新", use_container_width=True, key="manual_refresh"):
            st.cache_data.clear()
            for k in ['scan_df', 'buy_df', 'scan_time']:
                st.session_state.pop(k, None)
            st.rerun()

    # ==================== 初始化 AI 引擎 ====================
    engine = AIEngine(use_llm=use_llm, llm_provider=llm_provider, api_key=api_key)

    # ==================== 获取数据 ====================
    with st.spinner("⏳ 正在获取数据..."):
        if use_real_data:
            df = cached_get_kline(selected_code, analysis_days)
            if df.empty or len(df) < 30:
                st.warning(f"⚠️ {selected_code} 真实数据获取失败或数据不足，切换为模拟数据")
                use_real_data = False
        else:
            df = pd.DataFrame()

        if not use_real_data or df.empty:
            df = _generate_mock_kline(selected_code, analysis_days, stock_info)

        # 计算技术指标
        if show_all_ta:
            ti = TechnicalIndicators(df)
            df = ti.add_all().result()
        else:
            from technical_indicators import add_indicators_streamlined
            df = add_indicators_streamlined(df)

        # 并行获取其他数据
        fund_flow = cached_get_fund_flow(selected_code) if use_real_data else _mock_fund_flow()
        hotspots = cached_get_sector_hotspots() if use_real_data else _mock_hotspots()
        us_link = cached_get_us_linkage() if use_real_data else _mock_us_link()
        macro = get_macro_prediction()
        market = get_market_overview() if use_real_data else _mock_market()
        limit_up = get_limit_up_data() if use_real_data else _mock_limit_up()

    # ==================== AI分析 ====================
    signals = engine.detect_signals(df)
    risk = engine.calculate_risk_score(df, signals)
    levels = engine.calculate_trade_levels(df)
    chip_data = TechnicalIndicators(df.copy()).get_chip_distribution()

    # ==================== 🐉 一键擒龙结果展示 ====================
    if st.session_state.get('dragon_scan'):
        st.markdown("---")
        st.header("🐉 一键擒龙 — 扫描结果")

        dragon_pool_name = st.session_state.get('dragon_pool', '全部板块')
        dragon_days_val = st.session_state.get('dragon_days', 90)
        dragon_workers_val = st.session_state.get('dragon_workers', 5)

        with st.spinner("⏳ 正在并发扫描全市场股票..."):
            progress_bar = st.progress(0)
            status_text = st.empty()

            all_stocks = []
            if dragon_pool_name == "全部板块":
                for sector, stocks in STOCK_POOLS.items():
                    for code, name in stocks:
                        all_stocks.append((code, name, sector))
                for s in load_watchlist():
                    all_stocks.append((s['code'], s.get('name',''), s.get('sector','用户关注')))
            elif dragon_pool_name == "⭐ 我的关注":
                for s in load_watchlist():
                    all_stocks.append((s['code'], s.get('name',''), s.get('sector','用户关注')))
            else:
                for code, name in STOCK_POOLS.get(dragon_pool_name, []):
                    all_stocks.append((code, name, dragon_pool_name))

            # 排除科创板
            if st.session_state.get('dragon_exclude_star', True):
                all_stocks = [(c, n, s) for c, n, s in all_stocks if not c.startswith('688')]

            total = len(all_stocks)
            status_text.text(f"共 {total} 只股票待扫描...")
            scanned = 0
            scan_results_raw = []

            from stock_scanner import _scan_single_stock
            from concurrent.futures import ThreadPoolExecutor, as_completed

            pool = st.session_state.get('dragon_pool', '全部板块')
            exclude_star_val = st.session_state.get('dragon_exclude_star', True)

            with ThreadPoolExecutor(max_workers=dragon_workers_val) as executor:
                futures = {}
                for code, name, sector in all_stocks:
                    if exclude_star_val and code.startswith('688'):
                        continue
                    f = executor.submit(_scan_single_stock, code, name, sector, dragon_days_val)
                    futures[f] = (code, name, sector)

                for f in as_completed(futures):
                    try:
                        r = f.result()
                        if r:
                            scan_results_raw.append(r)
                    except Exception:
                        pass
                    scanned += 1
                    progress_bar.progress(min(scanned / max(total, 1), 1.0))
                    status_text.text(f"扫描进度: {scanned}/{total} — 已发现 {len(scan_results_raw)} 只有效结果")

            scan_df = pd.DataFrame(scan_results_raw) if scan_results_raw else pd.DataFrame()
            if not scan_df.empty and 'buy_score' in scan_df.columns:
                scan_df = scan_df.sort_values('buy_score', ascending=False)

            progress_bar.progress(1.0)
            status_text.text(f"扫描完成! 共分析 {len(scan_df)} 只股票")

            valid_scan = not scan_df.empty and 'buy_score' in scan_df.columns and 'signal' in scan_df.columns
            if valid_scan:
                # 保存到 session_state 供龙头看板/资金趋势复用
                st.session_state['scan_df'] = scan_df

                buy_df = filter_buy_signals(scan_df, min_buy_score=10)

                # Metric cards
                mc1, mc2, mc3, mc4, mc5 = st.columns(5)
                strong = len(scan_df[scan_df['signal'] == 'STRONG_BUY']) if 'signal' in scan_df.columns else 0
                mc1.metric("🟢 强烈买入", strong)
                buy_cnt = len(scan_df[scan_df['signal'] == 'BUY']) if 'signal' in scan_df.columns else 0
                mc2.metric("🔵 买入", buy_cnt)
                weak = len(scan_df[scan_df['signal'] == 'WEAK_BUY']) if 'signal' in scan_df.columns else 0
                mc3.metric("⚪ 偏多", weak)
                if not scan_df.empty:
                    top_sec = scan_df.groupby('sector')['buy_score'].mean().idxmax()
                    mc4.metric("📊 最强板块", top_sec)
                else:
                    mc4.metric("📊 最强板块", "-")
                if not buy_df.empty:
                    mc5.metric("⭐ 最高分", f"{buy_df.iloc[0]['name']} ({buy_df.iloc[0]['buy_score']}分)")
                else:
                    mc5.metric("⭐ 最高分", "-")

                # Sector summary
                st.subheader("📊 行业扫描汇总")
                sector_summ = summarize_by_sector(scan_df)
                if not sector_summ.empty:
                    st.dataframe(
                        sector_summ.style.background_gradient(
                            subset=['买入信号数', '平均买入分'], cmap='RdYlGn'
                        ), use_container_width=True,
                    )

                # Detail table
                show_all = st.checkbox("显示全部股票（含卖出/中性信号）", value=True, key="dragon_show_all")
                st.subheader("🔍 扫描结果清单")
                if show_all:
                    display_df = scan_df.copy()
                else:
                    display_df = buy_df.copy() if not buy_df.empty else scan_df.copy()
                display_df = display_df.sort_values('buy_score', ascending=False)

                show_cols = ['code', 'name', 'sector', 'price', 'change%',
                              'rsi',
                              'signal', 'buy_score', 'sell_score',
                              'buy_strategy', 'sell_strategy',
                              'buy_reasons', 'sell_reasons',
                              'ret_20d', 'ret_60d', 'macd_cross_wr']
                avail = [c for c in show_cols if c in display_df.columns]

                # 翻译信号和列名
                display_renamed = display_df[avail].copy()
                if 'signal' in display_renamed.columns:
                    display_renamed['signal'] = display_renamed['signal'].map(SIGNAL_CN).fillna(display_renamed['signal'])
                if 'buy_reasons' in display_renamed.columns:
                    display_renamed['buy_reasons'] = display_renamed['buy_reasons'].apply(_label_reasons)
                display_renamed.rename(columns=COL_CN, inplace=True)

                def color_signal(val):
                    colors = {
                        'STRONG_BUY': 'background-color: #ff4d4d; color: white',
                        'BUY': 'background-color: #ff8080; color: white',
                        'WEAK_BUY': 'background-color: #ffcccc',
                        'NEUTRAL': '', 'WEAK_SELL': 'background-color: #ccffcc',
                        'SELL': 'background-color: #80ff80',
                        'STRONG_SELL': 'background-color: #00cc00; color: white',
                        # 中文版
                        '🟢强烈买入': 'background-color: #ff4d4d; color: white',
                        '🔵买入': 'background-color: #ff8080; color: white',
                        '⚪偏多': 'background-color: #ffcccc',
                        '⬜中性': '', '🟡偏空': 'background-color: #ccffcc',
                        '🟠卖出': 'background-color: #80ff80',
                        '🔴强烈卖出': 'background-color: #00cc00; color: white',
                    }
                    return colors.get(val, '')

                def color_score_grad(val):
                    """buy_score颜色渐变"""
                    if val is None or pd.isna(val): return ''
                    v = float(val)
                    if v >= 50: return 'background-color: #d50000; color: white'
                    if v >= 35: return 'background-color: #ff4d4d; color: white'
                    if v >= 20: return 'background-color: #ff8080'
                    if v >= 10: return 'background-color: #ffcccc'
                    return ''

                try:
                    styled = display_renamed.style.applymap(color_signal, subset=[COL_CN['signal']]).applymap(
                        color_score_grad, subset=[COL_CN['buy_score']]).format({
                        COL_CN['price']: '{:.2f}', COL_CN['change%']: '{:+.2f}%', COL_CN['rsi']: '{:.1f}',
                        COL_CN['ret_20d']: '{:+.1f}%', COL_CN['ret_60d']: '{:+.1f}%',
                        COL_CN['macd_cross_wr']: '{:.1f}%',
                    })
                    st.dataframe(styled, use_container_width=True, height=500)
                except Exception:
                    st.warning("⚠️ 渲染表格出错，尝试纯文本显示")
                    st.dataframe(display_renamed, use_container_width=True, height=500)

                csv_out = display_df.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    "⬇️ 导出CSV", csv_out,
                    f"一键擒龙_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    "text/csv", use_container_width=True,
                    key="dl_dragon_sidebar",
                )

                # === 紧凑回测 ===
                with st.expander("🧪 快速回测验证", expanded=False):
                    from backtest_engine import BUILTIN_STRATEGIES, run_backtest
                    from technical_indicators import add_indicators_streamlined

                    col_sb1, col_sb2 = st.columns(2)
                    with col_sb1:
                        bt_strat_sb = st.selectbox("策略", list(BUILTIN_STRATEGIES.keys()),
                                                   key="bt_strat_sidebar")
                        desc_map = {
                            'MACD金叉+放量': 'MACD金叉且放量买入，适合趋势股',
                            '均线系统': '均线多头排列时买入，跟随趋势',
                            '布林带反转': '超跌反弹抄底，适合震荡市',
                            '多因子综合': '六维信号综合打分，最稳健',
                            '突破放量': '突破前高追涨，适合强势股',
                        }
                        st.caption(desc_map.get(bt_strat_sb, ''))
                    with col_sb2:
                        top_n_sb = st.selectbox("回测前N只", [3, 5, 8, 10, 15],
                                                index=1, key="bt_topn_sidebar")

                    if st.button("⚡ 快速回测", use_container_width=True, key="bt_run_sidebar"):
                        top_codes = scan_df.head(top_n_sb)['code'].tolist()
                        bt_sb_results = []
                        prog = st.progress(0)
                        for i, code in enumerate(top_codes):
                            try:
                                kdf = cached_get_kline(code, days=250)
                                if not kdf.empty and len(kdf) >= 60:
                                    kdf = add_indicators_streamlined(kdf.copy())
                                    r = run_backtest(kdf, bt_strat_sb, 100000)
                                    row = scan_df[scan_df['code'] == code].iloc[0]
                                    r['code'] = code
                                    r['name'] = row['name']
                                    r['sector'] = row['sector']
                                    r['buy_score'] = row['buy_score']
                                    bt_sb_results.append(r)
                            except Exception:
                                pass
                            prog.progress((i + 1) / len(top_codes))

                        if bt_sb_results:
                            st.success(f"回测完成 {len(bt_sb_results)} 只")
                            bt_tbl = []
                            for r in sorted(bt_sb_results, key=lambda x: x.get('total_return', -999), reverse=True):
                                trades = r.get('trades_count', 0) or 0
                                no_sig = (trades == 0)
                                bt_tbl.append({
                                    '代码': r['code'], '名称': r['name'],
                                    '收益%': '无信号' if no_sig else f"{r.get('total_return', 0):+.1f}%",
                                    '胜率%': '无信号' if no_sig else f"{r.get('win_rate', 0):.0f}%",
                                    '回撤%': '无信号' if no_sig else f"{r.get('max_drawdown', 0):.1f}%",
                                    '交易': 0 if no_sig else trades,
                                })
                            st.dataframe(pd.DataFrame(bt_tbl), use_container_width=True, height=250)
                            nosig_count = sum(1 for r in bt_sb_results if (r.get('trades_count', 0) or 0) == 0)
                            if nosig_count > 0:
                                st.caption(f"🟠 {nosig_count} 只股票无信号（该策略在历史数据中未触发买卖条件），可切换策略试试")
                        else:
                            st.warning("回测无有效结果")

                if st.button("🗑️ 关闭结果"):
                    st.session_state['dragon_scan'] = False
                    st.rerun()
            else:
                st.warning("扫描无结果，请重试")
                if st.button("🗑️ 关闭"):
                    st.session_state['dragon_scan'] = False
                    st.rerun()

        st.markdown("---")

    # ==================== 主Tab ====================
    tabs = st.tabs([
        "📊 技术分析",
        "🔥 热点扫描",
        "💼 持仓优化",
        "📉 行情复盘",
        "🔬 价值研报",
        "🧪 策略回测",
        "🐉 一键擒龙",
        "🏆 龙头看板",
        "💰 资金趋势",
    ])

    # ==================== TAB 1: 热点扫描 ====================
    with tabs[1]:
        st.header("当前市场热点 & 板块轮动")

        col1, col2 = st.columns([3, 2])

        with col1:
            st.subheader("📊 板块热度排行")
            display_cols = ['板块名称', '今日涨幅%', '5日涨幅%', '20日涨幅%', '主力净流入(亿)', '轮动预测']
            available = [c for c in display_cols if c in hotspots.columns]
            st.dataframe(
                hotspots[available].style.background_gradient(
                    subset=[c for c in ['今日涨幅%', '主力净流入(亿)'] if c in available],
                    cmap='RdYlGn',
                ),
                use_container_width=True, height=400,
            )

            st.subheader("🇺🇸 美股联动 (昨夜涨跌 → A股影响)")
            if us_link:
                us_rows = []
                for tk, info in us_link.items():
                    us_rows.append({
                        '代码': tk,
                        '昨夜涨跌%': info.get('yesterday_change', 0),
                        '关联A股': info.get('a_stock_related', ''),
                        '影响': info.get('impact', ''),
                    })
                us_df = pd.DataFrame(us_rows)
                st.dataframe(us_df.style.background_gradient(subset=['昨夜涨跌%'], cmap='RdYlGn'),
                             use_container_width=True)
            else:
                st.info("美股数据暂不可用")

        with col2:
            st.subheader("🏛️ 宏观预测")
            for k, v in macro.items():
                with st.expander(k):
                    st.write(v)

            st.subheader("📈 三大指数")
            if market:
                for name, info in market.items():
                    chg = info.get('change_pct', 0)
                    color = "red" if chg > 0 else "green" if chg < 0 else "gray"
                    st.metric(name, info.get('price', '-'), f"{chg:+.2f}%" if isinstance(chg, (int, float)) else str(chg))

            st.subheader("🎯 轮动预测")
            rotation = engine.predict_sector_rotation(hotspots)
            st.info(f"**市场判断**: {rotation['rotation_signal']}")
            if rotation.get('next_opportunity'):
                st.success("**下一阶段机会**:")
                for opp in rotation['next_opportunity']:
                    st.write(f"- {opp['name']}: {opp['reason']}")

            st.subheader("📰 相关新闻")
            news_list = get_news()
            for n in news_list:
                sentiment_emoji = "🟢" if n['sentiment'] == '正面' else "🔴" if n['sentiment'] == '负面' else "⚪"
                st.caption(f"{sentiment_emoji} [{n['time']}] {n['title']}")

        st.subheader("🔍 一键筛选")
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            if st.button("低位+资金流入板块", use_container_width=True):
                candidates = hotspots[
                    (hotspots['低位信号'].str.contains('强|中', na=False)) &
                    (hotspots['主力净流入(亿)'] > 3)
                    ] if '低位信号' in hotspots.columns else hotspots.head(5)
                st.success(f"筛选出 {len(candidates)} 个候选板块")
                st.dataframe(candidates)
        with col_b:
            if st.button("趋势强势板块", use_container_width=True):
                candidates = hotspots[hotspots['20日涨幅%'] > 5] if '20日涨幅%' in hotspots.columns else hotspots.head(5)
                st.dataframe(candidates)
        with col_c:
            if st.button("美股共振板块", use_container_width=True):
                st.info("美股半导体普涨 → A股半导体/存储/设备板块大概率高开")

    # ==================== TAB 0: 技术分析（默认首页） ====================
    with tabs[0]:
        st.header(f"📊 {selected_code} {stock_info['name']} — 技术全景分析")

        # K线主图
        st.plotly_chart(plot_kline(df, f"{stock_info['name']}({selected_code})"), use_container_width=True)

        # 信号仪表 + 关键指标卡片
        col1, col2, col3, col4, col5 = st.columns(5)
        latest = df.iloc[-1]
        prev = df.iloc[-2]
        change = (latest['close'] - prev['close']) / prev['close'] * 100

        col1.metric("最新价", f"{latest['close']:.2f}", f"{change:+.2f}%")
        col2.metric("MACD", f"{latest.get('MACD', 0):.3f}",
                    "金叉" if latest['MACD'] > latest['Signal'] else "死叉")
        col3.metric("RSI(14)", f"{latest.get('RSI', 0):.1f}",
                    "超买" if latest.get('RSI', 0) > 70 else "超卖" if latest.get('RSI', 0) < 30 else "中性")
        col4.metric("ADX", f"{latest.get('ADX', 0):.1f}",
                    "强趋势" if latest.get('ADX', 0) > 40 else "趋势" if latest.get('ADX', 0) > 25 else "震荡")
        col5.metric("量比", f"{latest.get('VOLUME_RATIO', 1):.1f}x",
                    "放量" if latest.get('VOLUME_RATIO', 1) > 1.5 else "正常")

        # 第二行指标
        st.markdown("---")
        st.subheader("📐 多维度技术指标一览")

        ta_cols = st.columns(6)
        indicator_data = [
            ("MA5", latest.get('MA5', 0), "MA20", latest.get('MA20', 0)),
            ("K", latest.get('K', 0), "D", latest.get('D', 0)),
            ("J", latest.get('J', 0), "WR", latest.get('WR', 0)),
            ("OBV", f"{latest.get('OBV', 0)/1e6:.1f}M", "CCI", latest.get('CCI', 0)),
            ("BB上轨", latest.get('BB_UPPER', 0), "BB下轨", latest.get('BB_LOWER', 0)),
            ("波动率", f"{latest.get('VOLATILITY', 0)*100:.1f}%", "夏普", f"{latest.get('SHARPE', 0):.2f}"),
        ]
        for i, ((label1, val1, label2, val2)) in enumerate(indicator_data):
            with ta_cols[i]:
                st.metric(label1, f"{val1:.2f}" if isinstance(val1, float) else val1)
                st.metric(label2, f"{val2:.2f}" if isinstance(val2, float) else val2)

        # 买卖价位
        st.markdown("---")
        st.subheader("🎯 买卖价位 & 止损止盈")
        buy_col, sell_col, stop_col, hold_col = st.columns(4)
        with buy_col:
            st.warning(f"**激进买入**: {levels['buy_point_aggressive']}")
            st.warning(f"**保守买入**: {levels['buy_point_conservative']}")
            st.warning(f"**抄底价**: {levels['buy_point_bottom_fishing']}")
        with sell_col:
            st.success(f"**止盈1**: {levels['take_profit_1']} (+{((levels['take_profit_1']-levels['current_price'])/levels['current_price']*100):.1f}%)")
            st.success(f"**止盈2**: {levels['take_profit_2']}")
            st.success(f"**长线目标**: {levels['take_profit_3']}")
        with stop_col:
            st.error(f"**紧止损**: {levels['stop_loss_tight']} (-{((levels['current_price']-levels['stop_loss_tight'])/levels['current_price']*100):.1f}%)")
            st.error(f"**正常止损**: {levels['stop_loss_normal']}")
            st.error(f"**最大止损**: {levels['stop_loss_wide']}")
        with hold_col:
            st.info(f"**持仓周期**: {levels['hold_period']}")
            st.info(f"**ATR**: {levels['atr']}")
            st.info(f"**MA60**: {levels['ma60']}")

        # 筹码分布
        st.markdown("---")
        col_ch1, col_ch2 = st.columns([2, 1])
        with col_ch1:
            plot_chip(df)
        with col_ch2:
            st.subheader("📊 筹码分析")
            if chip_data:
                st.metric("筹码密集区", f"{chip_data['max_chip_price']:.2f}")
                st.metric("平均成本", f"{chip_data['avg_cost']:.2f}")
                st.metric("获利比例", f"{chip_data['profit_ratio']:.1f}%")
                conc = chip_data['cost_std'] / chip_data['avg_cost'] * 100 if chip_data['avg_cost'] > 0 else 0
                st.metric("集中度", f"{conc:.1f}%",
                          "高" if conc < 10 else "中" if conc < 20 else "分散")

        # 实时信号面板
        st.markdown("---")
        st.subheader("🤖 AI综合信号")

        col_sig, col_risk = st.columns([1, 1])
        with col_sig:
            signal_map = {
                'STRONG_BUY': ('🟢🟢 强烈买入', 'signal-buy'),
                'BUY': ('🟢 买入', 'signal-buy'),
                'WEAK_BUY': ('🔵 偏多', ''),
                'NEUTRAL': ('⚪ 中性观望', ''),
                'WEAK_SELL': ('🟡 偏空', ''),
                'SELL': ('🔴 卖出', 'signal-sell'),
                'STRONG_SELL': ('🔴🔴 强烈卖出', 'signal-sell'),
            }
            label, cls = signal_map.get(signals['overall'], ('未知', ''))
            st.markdown(f"### 综合信号: <span class='{cls}'>{label}</span>", unsafe_allow_html=True)

            if signals['buy_signals']:
                st.write("**✅ 买入信号**:")
                for s in signals['buy_signals']:
                    st.write(f"- {s}")
            if signals['sell_signals']:
                st.write("**❌ 卖出信号**:")
                for s in signals['sell_signals']:
                    st.write(f"- {s}")
            if signals['neutral_signals']:
                st.write("**⚠️ 关注**:")
                for s in signals['neutral_signals']:
                    st.write(f"- {s}")

        with col_risk:
            plot_signal_gauge(signals['buy_score'], signals['sell_score'])

            st.markdown(f"### 风险评级: <span class='risk-{['high','mid','low'][0 if risk['score']>=70 else 1 if risk['score']>=40 else 2]}'>{risk['level']}</span>",
                        unsafe_allow_html=True)
            st.progress(risk['score'] / 100, text=f"{risk['score']}/100")
            if risk['factors']:
                st.write("**风险因素**:")
                for f in risk['factors']:
                    st.write(f"- {f}")

    # ==================== TAB 3: 持仓优化 ====================
    with tabs[2]:
        st.header("💼 持仓综合诊断 & 组合优化")

        total_capital = total_capital_input * 10000
        holdings_list = []
        for _, row in edited.iterrows():
            code = str(row['code']).zfill(6)
            current_price = row.get('cost', 100) * 1.05
            try:
                fund_f = cached_get_fund_flow(code) if use_real_data else _mock_fund_flow()
                current_price = fund_f.get('最新价', current_price)
            except Exception:
                pass

            holdings_list.append({
                'code': code,
                'name': row.get('name', ''),
                'cost': float(row['cost']),
                'shares': int(row['shares']),
                'position': float(row['cost']) * int(row['shares']),
                'position_pct': float(row['cost']) * int(row['shares']) / total_capital * 100,
                'current_price': current_price,
                'sector': row.get('sector', '未分类'),
            })

        st.subheader("📋 持仓清单")
        display_holdings = []
        for h in holdings_list:
            pnl = (h['current_price'] - h['cost']) / h['cost'] * 100
            display_holdings.append({
                '代码': h['code'],
                '名称': h['name'],
                '成本价': h['cost'],
                '现价': h['current_price'],
                '股数': h['shares'],
                '市值(万元)': round(h['position'] / 10000, 2),
                '占比': f"{h['position_pct']:.1f}%",
                '浮盈%': f"{pnl:+.2f}%",
            })
        st.dataframe(pd.DataFrame(display_holdings), use_container_width=True)

        st.divider()

        # 逐诊断
        st.subheader("🔍 逐个诊断")
        for h in holdings_list:
            code = h['code']
            with st.expander(f"{code} {h['name']} — 成本 {h['cost']:.2f} | 现价 {h['current_price']:.2f} | 浮盈 {((h['current_price']-h['cost'])/h['cost']*100):+.2f}%"):
                try:
                    stock_df = cached_get_kline(code, 120) if use_real_data else None
                    if stock_df is None or stock_df.empty:
                        stock_df = _generate_mock_kline(code, 120, STOCK_POOL.get(code, stock_info))
                    stock_df = add_technical_indicators(stock_df)

                    fund_f = cached_get_fund_flow(code) if use_real_data else _mock_fund_flow()
                    diag = engine.diagnose_holding(
                        code, h['name'], h['cost'], h['current_price'],
                        h['position_pct'], stock_df, fund_f,
                    )

                    st.markdown(f"### 诊断结果: **{diag['action']}**")
                    st.info(diag['advice'])
                    for d in diag['details']:
                        st.write(f"- {d}")

                    col_d1, col_d2 = st.columns(2)
                    with col_d1:
                        st.metric("信号", diag['signal'])
                    with col_d2:
                        st.metric("风险", f"{diag['risk']['level']} ({diag['risk']['score']}/100)")

                except Exception as e:
                    st.error(f"诊断异常: {e}")

        st.divider()

        # 组合建议
        st.subheader("📊 组合配置建议")
        portfolio = engine.portfolio_advice(holdings_list, total_capital)
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            st.write("**当前配置**")
            st.metric("总仓位", f"{100 - portfolio['cash_pct']:.1f}%")
            st.metric("现金占比", f"{portfolio['cash_pct']:.1f}%")
            st.write("**建议配置**")
            st.info(f"🟢 核心仓位 60%: {portfolio['suggested_allocation']['core']['desc']}")
            st.info(f"🔵 卫星仓位 30%: {portfolio['suggested_allocation']['satellite']['desc']}")
            st.warning(f"🟡 防御仓位 10%: {portfolio['suggested_allocation']['defensive']['desc']}")
        with col_p2:
            st.write("**风控检查**")
            if portfolio['warnings']:
                for w in portfolio['warnings']:
                    st.warning(w)
            if portfolio['actions']:
                for a in portfolio['actions']:
                    st.info(f"💡 {a}")

    # ==================== TAB 4: 行情复盘 ====================
    with tabs[3]:
        st.header("📉 大盘走势 & 行情复盘")

        col_r1, col_r2 = st.columns([2, 1])

        with col_r1:
            st.subheader("今日市场风格")
            sentiment = engine.market_sentiment(
                df, limit_up.get('limit_up', 0),
                limit_up.get('limit_down', 0),
            )
            st.markdown(f"### 总体判断: **{sentiment['overall']}**")
            st.markdown(f"**风格**: {sentiment['market_style']} | **风险偏好**: {sentiment['risk_appetite']}")
            for d in sentiment.get('details', []):
                st.write(f"- {d}")

            st.subheader("资金动态")
            fund_data = pd.DataFrame({
                '指标': ['主力净流入', '北向资金', '涨停板', '跌停板', 'ETF申赎'],
                '数值': [
                    f"主力净额：+{hotspots['主力净流入(亿)'].sum() if '主力净流入(亿)' in hotspots.columns else 0:.0f}亿",
                    '北向：+20.8亿(预估)',
                    f'{limit_up.get("limit_up", 0)}家',
                    f'{limit_up.get("limit_down", 0)}家',
                    '科技ETF净申购 12.5亿',
                ],
            })
            st.dataframe(fund_data, use_container_width=True, hide_index=True)

            st.subheader("板块轮动复盘")
            if not hotspots.empty:
                hot = hotspots.nlargest(5, '今日涨幅%') if '今日涨幅%' in hotspots.columns else hotspots.head(5)
                cold = hotspots.nsmallest(5, '今日涨幅%') if '今日涨幅%' in hotspots.columns else hotspots.tail(5)
                col_h, col_c = st.columns(2)
                with col_h:
                    st.write("**🔥 今日领涨**")
                    st.dataframe(hot[['板块名称', '今日涨幅%', '主力净流入(亿)']], use_container_width=True)
                with col_c:
                    st.write("**❄️ 今日领跌**")
                    st.dataframe(cold[['板块名称', '今日涨幅%', '主力净流入(亿)']], use_container_width=True)

        with col_r2:
            st.subheader("📌 下一阶段机会")
            rotation = engine.predict_sector_rotation(hotspots)
            st.success(f"**方向**: {rotation['rotation_signal']}")

            st.write("**高概率方向**:")
            st.write("- 🟢 存储芯片 (低位+美股MU联动+资金流入)")
            st.write("- 🟢 半导体设备 (国产替代加速+订单增加)")
            st.write("- 🔵 AI算力 (英伟达股价传导+需求旺盛)")
            st.write("**规避方向**:")
            st.write("- 🔴 高估值白酒/消费 (缺乏增量资金)")
            st.write("- 🟡 银行 (轮动资金外流)")

            st.subheader("⚠️ 今日预警")
            if limit_up.get('limit_down', 0) > 30:
                st.error(f"跌停家数 {limit_up['limit_down']} 家，警惕系统性风险")
            if '主力净流入(亿)' in hotspots.columns and hotspots['主力净流入(亿)'].sum() < -50:
                st.warning("主力资金总体流出，控制仓位")
            st.info("关注北向资金持续流入板块")

    # ==================== TAB 5: 价值研报 ====================
    with tabs[4]:
        st.header(f"🔬 {selected_code} {stock_info['name']} — 深度价值研究报告")

        # 财务概要
        st.subheader("📋 基本面概览")
        fin = get_financial_summary(selected_code) if use_real_data else _mock_financial()
        if fin:
            cols = st.columns(5)
            cols[0].metric("最新价", fin.get('最新价', '-'))
            cols[1].metric("PE(TTM)", f"{fin.get('PE(TTM)', 0):.1f}")
            cols[2].metric("PB", f"{fin.get('PB', 0):.1f}")
            cols[3].metric("总市值(亿)", fin.get('总市值(亿)', '-'))
            cols[4].metric("52周高", fin.get('52周最高', '-'))

        # 研报正文
        report = engine.generate_report(
            selected_code, stock_info['name'], stock_info['sector'],
            df, signals, risk, levels, fund_flow, chip_data,
            us_link if us_link else None, macro,
        )
        st.markdown(report)

        # 导出报告
        col_export1, col_export2 = st.columns(2)
        with col_export1:
            if st.button("📄 导出 DOCX 研报", use_container_width=True):
                buffer = create_docx_report(report)
                st.download_button(
                    "⬇️ 下载DOCX研报",
                    buffer,
                    f"{selected_code}_研究报告_{datetime.now().strftime('%Y%m%d')}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key="dl_docx",
                )
        with col_export2:
            st.download_button(
                "⬇️ 下载MD研报",
                report,
                f"{selected_code}_研究报告_{datetime.now().strftime('%Y%m%d')}.md",
                "text/markdown",
                use_container_width=True, key="dl_md",
            )

        # LLM解读
        if use_llm:
            st.subheader("🤖 AI大模型解读")
            with st.spinner("AI分析中..."):
                enhanced = engine.enhance_with_llm(signals)
                st.info(enhanced)

    # ==================== TAB 6: 策略回测 ====================
    with tabs[5]:
        st.header("🧪 策略回测实验室")

        st.info("回测说明：以下回测基于所选股票的历史K线数据。真实项目建议使用 vectorbt 或 backtrader 做更完整的回测。")

        col_bt1, col_bt2 = st.columns([1, 1])

        with col_bt1:
            st.subheader("📊 单策略回测")
            bt_strategy = st.selectbox(
                "选择策略",
                list(BUILTIN_STRATEGIES.keys()),
                help="选择内置经典策略进行回测",
            )
            bt_capital = st.number_input("初始资金", 10000, 10000000, 100000, step=10000)

            if st.button("▶️ 执行回测", use_container_width=True):
                with st.spinner("回测运行中..."):
                    result = run_backtest(df, bt_strategy, bt_capital)
                    if 'error' in result:
                        st.error(result['error'])
                    else:
                        st.success(f"回测完成 — 总收益: {result['total_return']:.2f}%")

                        col_r1, col_r2, col_r3, col_r4 = st.columns(4)
                        col_r1.metric("总收益率", f"{result['total_return']:.2f}%")
                        col_r2.metric("年化收益", f"{result['annual_return']:.2f}%")
                        col_r3.metric("最大回撤", f"{result['max_drawdown']:.2f}%")
                        col_r4.metric("胜率", f"{result['win_rate']:.1f}%")

                        col_r5, col_r6, col_r7 = st.columns(3)
                        col_r5.metric("夏普比率", f"{result['sharpe_ratio']:.3f}")
                        col_r6.metric("交易次数", result['trades_count'])
                        col_r7.metric("最终资金", f"{result['final_value']:,.0f}")

                        # 权益曲线
                        if 'equity_curve' in result and not result['equity_curve'].empty:
                            eq = result['equity_curve']
                            fig_eq = go.Figure()
                            fig_eq.add_trace(go.Scatter(
                                x=eq['date'], y=eq['equity'],
                                name='权益曲线', fill='tozeroy',
                                fillcolor='rgba(30,144,255,0.1)',
                                line=dict(color='#1f77b4'),
                            ))
                            fig_eq.update_layout(
                                title="权益曲线", height=350,
                                xaxis_title="日期", yaxis_title="账户权益(元)",
                                margin=dict(l=0, r=0, t=40, b=0),
                            )
                            st.plotly_chart(fig_eq, use_container_width=True)

        with col_bt2:
            st.subheader("📈 策略对比")
            if st.button("⚔️ 对比所有策略", use_container_width=True):
                with st.spinner("多策略对比中..."):
                    comp = compare_strategies(df)
                    st.dataframe(
                        comp.style.background_gradient(
                            subset=['total_return', 'annual_return', 'sharpe_ratio'],
                            cmap='RdYlGn',
                        ).highlight_min(subset=['max_drawdown'], color='lightgreen'),
                        use_container_width=True,
                    )

        st.divider()

        # 策略说明
        st.subheader("📚 策略库")
        for name, info in BUILTIN_STRATEGIES.items():
            with st.expander(f"{info['type']} — {name}"):
                st.write(f"**描述**: {info['desc']}")
                st.write(f"**类型**: {info['type']}")

        # 自定义策略
        st.subheader("⚙️ 自定义策略参数")
        with st.expander("高级参数设置"):
            st.write("生产环境支持更灵活的参数调优。当前Demo版本展示了回测框架的核心逻辑。")
            st.code("""
# 自定义策略示例
def my_strategy(df, idx, params):
    row = df.iloc[idx]
    if row['MACD_GOLDEN'] == 1 and row['VOLUME_RATIO'] > 2.0:
        return 'BUY'
    if row['RSI'] > 75:
        return 'SELL'
    return 'HOLD'
            """, language="python")

    # ==================== TAB 7: 一键擒龙 ====================
    with tabs[6]:
        st.header("🐉 一键擒龙 — 全市场智能扫描")

        st.markdown("""
        一键扫描A股全市场或指定板块，AI自动分析每只股票的技术面，筛选出**具有买入信号的股票**，
        按得分排序，帮你快速锁定值得关注的标的。
        """)

        col_ctrl1, col_ctrl2, col_ctrl3 = st.columns([1, 1, 1])

        with col_ctrl1:
            scan_pool_name = st.selectbox(
                "选择扫描范围",
                options=["全部板块"] + list(STOCK_POOLS.keys()),
                index=0,
                help="选择预设板块或扫描全部"
            )

        with col_ctrl2:
            scan_days = st.slider("分析天数", 30, 250, 90,
                                  help="K线数据量越多，分析越准确，但扫描越慢")
            scan_workers = st.slider("并发线程", 1, 10, 5,
                                     help="线程越多越快，但可能触发限流")

        with col_ctrl3:
            min_buy_score = st.slider("最低买入分数", 0, 80, 10,
                                      help="只显示买入评分≥此值的股票")
            sort_by = st.selectbox("排序方式", ["买入评分降序", "涨幅降序", "行业分组"])

        col_ctrl4, _, _ = st.columns([1, 1, 1])
        with col_ctrl4:
            exclude_star_tab = st.checkbox("排除科创板(688开头)", value=True, key="tab_exclude_star",
                                           help="科创板需要50万门槛，无权限者建议勾选")

        st.divider()

        if st.button("🚀 开始扫描", type="primary", use_container_width=True):
            scan_results = st.empty()

            with st.spinner("⏳ 正在扫描全市场股票，请稍候..."):
                progress_bar = st.progress(0)
                status_text = st.empty()

                try:
                    all_stocks = []
                    if scan_pool_name == "全部板块":
                        for sector, stocks in STOCK_POOLS.items():
                            for code, name in stocks:
                                all_stocks.append((code, name, sector))
                    else:
                        for code, name in STOCK_POOLS.get(scan_pool_name, []):
                            all_stocks.append((code, name, scan_pool_name))

                    # 排除科创板
                    if exclude_star_tab:
                        all_stocks = [(c, n, s) for c, n, s in all_stocks if not c.startswith('688')]

                    total_stocks = len(all_stocks)
                    status_text.text(f"共 {total_stocks} 只股票待扫描...")
                    scanned_tab = 0
                    results_tab = []

                    from stock_scanner import _scan_single_stock
                    from concurrent.futures import ThreadPoolExecutor, as_completed

                    pool_name = scan_pool_name

                    with ThreadPoolExecutor(max_workers=scan_workers) as executor:
                        futures_tab = {}
                        for code, name, sector in all_stocks:
                            f = executor.submit(_scan_single_stock, code, name, sector, scan_days)
                            futures_tab[f] = (code, name, sector)

                        for f in as_completed(futures_tab):
                            try:
                                r = f.result()
                                if r:
                                    results_tab.append(r)
                            except Exception:
                                pass
                            scanned_tab += 1
                            progress_bar.progress(min(scanned_tab / max(total_stocks, 1), 1.0))
                            status_text.text(f"扫描进度: {scanned_tab}/{total_stocks} — 已发现 {len(results_tab)} 只有效结果")

                    scan_df = pd.DataFrame(results_tab) if results_tab else pd.DataFrame()
                    if not scan_df.empty:
                        scan_df = scan_df.sort_values('buy_score', ascending=False)

                    scan_df_len = len(scan_df)
                    progress_bar.progress(1.0)
                    status_text.text(f"扫描完成! 共分析 {scan_df_len} 只股票")

                    if not scan_df.empty and 'buy_score' in scan_df.columns:
                        buy_df = filter_buy_signals(scan_df, min_buy_score=min_buy_score)
                        st.session_state['scan_df'] = scan_df
                        st.session_state['buy_df'] = buy_df
                        st.session_state['scan_time'] = datetime.now().strftime('%H:%M:%S')

                except Exception as e:
                    st.error(f"扫描异常: {e}")

        # Display cached results
        if 'scan_df' in st.session_state and not st.session_state['scan_df'].empty:
            scan_df = st.session_state['scan_df']
            buy_df = st.session_state.get('buy_df', pd.DataFrame())
            scan_time = st.session_state.get('scan_time', '')

            st.success(f"✅ 扫描完成 ({scan_time}) | 共 {len(scan_df)} 只 | "
                       f"买入 {len(buy_df)} 只 | 卖出/中性 {len(scan_df) - len(buy_df)} 只")

            st.divider()

            # Summary metrics
            m1, m2, m3, m4, m5 = st.columns(5)
            with m1:
                strong_buy = len(scan_df[scan_df['signal'] == 'STRONG_BUY']) if 'signal' in scan_df.columns else 0
                st.metric("强烈买入", strong_buy)
            with m2:
                buy = len(scan_df[scan_df['signal'] == 'BUY']) if 'signal' in scan_df.columns else 0
                st.metric("买入", buy)
            with m3:
                weak_buy = len(scan_df[scan_df['signal'] == 'WEAK_BUY']) if 'signal' in scan_df.columns else 0
                st.metric("偏多", weak_buy)
            with m4:
                top_sector = scan_df.groupby('sector')['buy_score'].mean().idxmax() if not scan_df.empty else '-'
                st.metric("最强板块", top_sector)
            with m5:
                if not buy_df.empty:
                    top_stock = buy_df.iloc[0] if sort_by != '涨幅降序' else buy_df.sort_values('change%', ascending=False).iloc[0]
                    st.metric("最高分个股", f"{top_stock['name']} ({top_stock['buy_score']}分)")

            # Sector summary
            st.subheader("📊 行业扫描汇总")
            sector_summary = summarize_by_sector(scan_df)
            if not sector_summary.empty:
                st.dataframe(
                    sector_summary.style.background_gradient(
                        subset=['买入信号数', '平均买入分'],
                        cmap='RdYlGn',
                    ),
                    use_container_width=True,
                )

            st.divider()

            # Detailed results
            st.subheader("🔍 详细扫描结果")
            show_all2 = st.checkbox("显示全部股票（含卖出/中性信号）", value=True, key="tab_show_all")
            st.markdown(f"**筛选条件**: 买入分数 ≥ {min_buy_score} | 排序: {sort_by} | {'全部信号' if show_all2 else '仅买入信号'}")

            if show_all2:
                display_df = scan_df.copy()
            else:
                display_df = buy_df.copy() if not buy_df.empty else scan_df.copy()

            if sort_by == '买入评分降序':
                display_df = display_df.sort_values('buy_score', ascending=False)
            elif sort_by == '涨幅降序':
                display_df = display_df.sort_values('change%', ascending=False)
            else:
                display_df = display_df.sort_values(['sector', 'buy_score'], ascending=[True, False])

            # Format for display
            show_cols = ['code', 'name', 'sector', 'price', 'change%',
                         'rsi',
                         'signal', 'buy_score', 'sell_score',
                         'buy_strategy', 'sell_strategy',
                         'buy_reasons', 'sell_reasons',
                         'ret_20d', 'ret_60d', 'macd_cross_wr']

            available_cols = [c for c in show_cols if c in display_df.columns]

            # 翻译信号和列名
            display_renamed = display_df[available_cols].copy()
            if 'signal' in display_renamed.columns:
                display_renamed['signal'] = display_renamed['signal'].map(SIGNAL_CN).fillna(display_renamed['signal'])
            if 'buy_reasons' in display_renamed.columns:
                display_renamed['buy_reasons'] = display_renamed['buy_reasons'].apply(_label_reasons)
            display_renamed.rename(columns=COL_CN, inplace=True)

            # Color coding for signal
            def color_signal2(val):
                colors = {
                    'STRONG_BUY': 'background-color: #ff4d4d; color: white',
                    'BUY': 'background-color: #ff8080; color: white',
                    'WEAK_BUY': 'background-color: #ffcccc',
                    'NEUTRAL': '',
                    'WEAK_SELL': 'background-color: #ccffcc',
                    'SELL': 'background-color: #80ff80',
                    'STRONG_SELL': 'background-color: #00cc00; color: white',
                    '🟢强烈买入': 'background-color: #ff4d4d; color: white',
                    '🔵买入': 'background-color: #ff8080; color: white',
                    '⚪偏多': 'background-color: #ffcccc',
                    '⬜中性': '', '🟡偏空': 'background-color: #ccffcc',
                    '🟠卖出': 'background-color: #80ff80',
                    '🔴强烈卖出': 'background-color: #00cc00; color: white',
                }
                return colors.get(val, '')

            def color_score_grad2(val):
                """buy_score颜色渐变"""
                if val is None or pd.isna(val): return ''
                v = float(val)
                if v >= 50: return 'background-color: #d50000; color: white'
                if v >= 35: return 'background-color: #ff4d4d; color: white'
                if v >= 20: return 'background-color: #ff8080'
                if v >= 10: return 'background-color: #ffcccc'
                return ''

            try:
                styled = display_renamed.style.applymap(
                    color_signal2, subset=[COL_CN['signal']]
                ).applymap(
                    color_score_grad2, subset=[COL_CN['buy_score']]
                ).format({
                    COL_CN['price']: '{:.2f}',
                    COL_CN['change%']: '{:+.2f}%',
                    COL_CN['rsi']: '{:.1f}',
                    COL_CN['ret_20d']: '{:+.1f}%',
                    COL_CN['ret_60d']: '{:+.1f}%',
                    COL_CN['macd_cross_wr']: '{:.1f}%',
                })
                st.dataframe(styled, use_container_width=True, height=500)
            except Exception:
                st.dataframe(display_renamed, use_container_width=True, height=500)

            # Stock detail expander
            with st.expander("📋 查看某只股票详细分析"):
                if not display_df.empty:
                    detail_code = st.selectbox(
                        "选择股票",
                        display_df['code'].tolist(),
                        format_func=lambda x: f"{x} {display_df[display_df['code']==x]['name'].iloc[0]}"
                    )
                    if detail_code:
                        detail_row = display_df[display_df['code'] == detail_code].iloc[0]
                        buy_strat = detail_row.get('buy_strategy', '')
                        sell_strat = detail_row.get('sell_strategy', '')
                        st.markdown(f"""
                        **{detail_row['code']} {detail_row['name']}**
                        - 板块: {detail_row['sector']} | 信号: {detail_row['signal']}
                        - 价格: {detail_row['price']} | 涨跌: {detail_row.get('change%', 0):+.2f}%
                        - 买入分: {detail_row['buy_score']} | 卖出分: {detail_row.get('sell_score', '-')}
                        - RSI: {detail_row['rsi']}
                        - 🟢 **买入策略**: {buy_strat}
                        - 🔴 **卖出策略**: {sell_strat}
                        - 20日涨幅: {detail_row.get('ret_20d', 0):+.1f}% | 60日涨幅: {detail_row.get('ret_60d', 0):+.1f}%
                        - 买入理由: {detail_row.get('buy_reasons', '')}
                        - 卖出理由: {detail_row.get('sell_reasons', '')}
                        """)
                else:
                    st.info("没有符合筛选条件的股票")

            # Export
            st.divider()

            # Per-stock watchlist checkboxes
            st.subheader("⭐ 逐只加入关注")
            st.caption("勾选股票后点击下方按钮批量加入关注列表")
            watch_candidates = display_df.copy()
            wl_cols = st.columns(4)
            stock_groups = []
            for i, (_, row) in enumerate(watch_candidates.iterrows()):
                group_idx = i % 4
                with wl_cols[group_idx]:
                    label = f"{row['code']} {row['name']}"
                    checked = is_in_watchlist(str(row['code']))
                    tick = st.checkbox(
                        f"{'✅' if checked else '⬜'} {label}",
                        value=checked,
                        key=f"wl_chk_{row['code']}",
                    )
                    if tick and not checked:
                        stock_groups.append(row)

            if stock_groups:
                if st.button(f"⭐ 加入勾选的 {len(stock_groups)} 只股票", use_container_width=True):
                    for r in stock_groups:
                        add_to_watchlist(str(r['code']), str(r['name']),
                                        str(r.get('sector', '')), float(r.get('price', 0)))
                    st.success(f"已添加 {len(stock_groups)} 只")
                    st.rerun()

            st.divider()
            export_col1, export_col2, export_col3 = st.columns(3)
            with export_col1:
                csv = display_df.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    "⬇️ 导出CSV", csv,
                    f"一键擒龙_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    "text/csv", use_container_width=True,
                )
            with export_col2:
                if st.button("⭐ 全部买入信号加入关注", use_container_width=True):
                    added = 0
                    bdf = display_df[display_df['signal'].isin(['STRONG_BUY','BUY','WEAK_BUY'])] if 'signal' in display_df.columns else display_df
                    for _, r in bdf.iterrows():
                        add_to_watchlist(str(r['code']), str(r['name']), str(r.get('sector','')),
                                        float(r.get('price', 0)))
                        added += 1
                    st.success(f"已添加 {added} 只")
                    st.rerun()
            with export_col3:
                if st.button("🔄 清空结果重新扫描", use_container_width=True):
                    for k in ['dragon_scan', 'scan_df', 'buy_df', 'scan_time']:
                        st.session_state.pop(k, None)
                    st.rerun()

                # ==== 🧪 策略回测验证 ====
            st.divider()
            with st.expander("🧪 一键擒龙策略回测验证", expanded=False):
                st.markdown("对扫描出的推荐股票，选用不同策略进行历史回测，验证信号质量。")

                from backtest_engine import BUILTIN_STRATEGIES, BacktestEngine, run_backtest
                from technical_indicators import add_indicators_streamlined

                # 策略白话说明
                STRATEGY_EXPLAIN = {
                    'MACD金叉+放量': '📖 当MACD出现金叉（快线上穿慢线）且成交量放大1.5倍时买入，MACD死叉或RSI超买(>75)时卖出。适合趋势明显的股票。',
                    '均线系统': '📖 当5日均线上穿20日均线且均线呈多头排列时买入，均线死叉或RSI>78时卖出。适合跟随中期趋势。',
                    '布林带反转': '📖 当股价跌破布林带下轨、RSI超卖(<35)且KDJ超卖时买入，突破上轨时卖出。适合震荡行情抄底逃顶。',
                    '多因子综合': '📖 综合MACD/KDJ/均线/RSI/WR/成交量六维信号打分，买入分≥5且卖出分≤1时买入，卖出分≥4时卖出。最稳健。',
                    '突破放量': '📖 当股价突破前期20日最高点且成交量放大2倍以上时追涨买入，MACD死叉或跌破20日均线时卖出。适合强势突破行情。',
                }

                strat_names = list(BUILTIN_STRATEGIES.keys())
                col_bt1, col_bt2, col_bt3 = st.columns(3)
                with col_bt1:
                    bt_strategy = st.selectbox("回测策略", strat_names, index=0, key="bt_strategy_tab7")
                    st.caption(STRATEGY_EXPLAIN.get(bt_strategy, ''))
                with col_bt2:
                    top_n = st.slider("回测前N只推荐股", 1, min(20, len(scan_df)), min(8, len(scan_df)), 1, key="bt_topn_tab7")
                with col_bt3:
                    bt_capital = st.number_input("初始资金(万)", 1, 1000, 10, 1, key="bt_capital_tab7")

                top_codes = scan_df.head(top_n)['code'].tolist()
                bt_selected = st.multiselect(
                    f"选择回测股票（默认前{top_n}只）",
                    options=scan_df['code'].tolist(),
                    default=top_codes,
                    format_func=lambda c: f"{c} {scan_df[scan_df['code']==c]['name'].iloc[0]}",
                    key="bt_stocks_tab7"
                )

                if st.button("🚀 开始批量回测", type="primary", use_container_width=True, key="bt_run_tab7"):
                    if not bt_selected:
                        st.warning("请选择至少一只股票")
                    else:
                        bt_results = []
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        total = len(bt_selected)

                        for i, code in enumerate(bt_selected):
                            status_text.text(f"回测中: {code} ({i+1}/{total})")
                            try:
                                kdf = cached_get_kline(code, days=250)
                                if kdf.empty or len(kdf) < 60:
                                    bt_results.append({'code': code, 'name': '数据不足', 'error': 'K线数据不足'})
                                else:
                                    kdf = add_indicators_streamlined(kdf.copy())
                                    result = run_backtest(kdf, bt_strategy, bt_capital * 10000)
                                    row_info = scan_df[scan_df['code'] == code].iloc[0]
                                    result['code'] = code
                                    result['name'] = row_info['name']
                                    result['sector'] = row_info['sector']
                                    result['buy_score'] = row_info['buy_score']
                                    result['signal'] = row_info['signal']
                                    bt_results.append(result)
                            except Exception as e:
                                bt_results.append({'code': code, 'name': '异常', 'error': str(e)[:50]})
                            progress_bar.progress((i + 1) / total)

                        progress_bar.progress(1.0)
                        status_text.text(f"回测完成! 共 {len(bt_results)} 只")
                        st.session_state['bt_results'] = bt_results
                        st.session_state['bt_strategy_name'] = bt_strategy

                if 'bt_results' in st.session_state and st.session_state['bt_results']:
                    bt_results = st.session_state['bt_results']
                    bt_name = st.session_state.get('bt_strategy_name', '')
                    st.markdown(f"**策略**: {bt_name} | 回测 {len(bt_results)} 只")

                    # Summary cards
                    valid = [r for r in bt_results if 'total_return' in r]
                    if valid:
                        mc1, mc2, mc3, mc4, mc5 = st.columns(5)
                        avg_ret = np.mean([r['total_return'] for r in valid])
                        mc1.metric("📈 平均收益", f"{avg_ret:+.1f}%")
                        avg_wr = np.mean([r.get('win_rate', 0) for r in valid])
                        mc2.metric("🎯 平均胜率", f"{avg_wr:.1f}%")
                        avg_dd = np.mean([r.get('max_drawdown', 0) for r in valid])
                        mc3.metric("📉 平均回撤", f"{avg_dd:.1f}%")
                        win_count = sum(1 for r in valid if r.get('total_return', 0) > 0)
                        mc4.metric("🏆 盈利数", f"{win_count}/{len(valid)}")
                        best = max(valid, key=lambda r: r.get('total_return', -999))
                        mc5.metric("⭐ 最佳", f"{best.get('name','')} ({best.get('total_return',0):+.1f}%)")

                    # Results table
                    bt_display = []
                    for r in bt_results:
                        trades = r.get('trades_count', 0) or 0
                        no_signal = (trades == 0 and 'error' not in r)
                        bt_display.append({
                            '代码': r.get('code', ''),
                            '名称': r.get('name', ''),
                            '板块': r.get('sector', ''),
                            'AI评分': r.get('buy_score', ''),
                            '总收益%': '无信号' if no_signal else r.get('total_return', '-'),
                            '年化%': '无信号' if no_signal else r.get('annual_return', '-'),
                            '最大回撤%': '无信号' if no_signal else r.get('max_drawdown', '-'),
                            '胜率%': '无信号' if no_signal else r.get('win_rate', '-'),
                            '夏普': '无信号' if no_signal else r.get('sharpe_ratio', '-'),
                            '交易次数': 0 if no_signal else trades,
                            '错误': r.get('error', ''),
                        })
                    bt_df_display = pd.DataFrame(bt_display)

                    def color_ret(val):
                        if isinstance(val, str): return ''
                        try: v = float(val); return 'background-color: #c8e6c9' if v > 0 else 'background-color: #ffcdd2'
                        except: return ''
                    def color_wr(val):
                        if isinstance(val, str): return ''
                        try: v = float(val); return 'background-color: #c8e6c9' if v > 50 else ''
                        except: return ''
                    def color_nosig(val):
                        return 'background-color: #fff3e0; color: #e65100' if val == '无信号' else ''

                    try:
                        styled_bt = bt_df_display.style \
                            .applymap(color_ret, subset=['总收益%', '年化%']) \
                            .applymap(color_wr, subset=['胜率%']) \
                            .applymap(color_nosig, subset=['总收益%'])
                        # 只格式化数值列
                        for col in ['总收益%', '年化%', '最大回撤%', '胜率%', '夏普']:
                            try:
                                bt_df_display[col] = pd.to_numeric(bt_df_display[col], errors='coerce')
                            except:
                                pass
                        st.dataframe(styled_bt, use_container_width=True, height=400)
                    except Exception:
                        st.dataframe(bt_df_display, use_container_width=True, height=400)

                    # Equity curve for best performer
                    if valid:
                        best_r = max(valid, key=lambda r: r.get('total_return', -999))
                        eq_df = best_r.get('equity_curve')
                        if eq_df is not None and not eq_df.empty:
                            st.subheader(f"📊 {best_r.get('name','')} 权益曲线")
                            fig_eq = go.Figure()
                            fig_eq.add_trace(go.Scatter(
                                x=eq_df['date'], y=eq_df['equity'],
                                name='权益曲线', fill='tozeroy',
                                fillcolor='rgba(30,144,255,0.1)',
                                line=dict(color='#1f77b4'),
                            ))
                            fig_eq.add_hline(y=bt_capital * 10000, line_dash="dash",
                                             line_color="gray", annotation_text="初始资金")
                            fig_eq.update_layout(height=300, margin=dict(l=0, r=0, t=10, b=0))
                            st.plotly_chart(fig_eq, use_container_width=True)

                    # Export
                    csv_bt = bt_df_display.to_csv(index=False).encode('utf-8-sig')
                    st.download_button("⬇️ 导出回测报告", csv_bt,
                                       f"回测_{bt_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                                       "text/csv", key="dl_bt_tab7")
        else:
            st.info("👆 点击上方 **开始扫描** 按钮，AI将自动分析全市场股票。")
            st.markdown("""
            **扫描范围预览**:
            """)
            preview_cols = st.columns(4)
            pool_names = list(STOCK_POOLS.keys())
            for i, name in enumerate(pool_names):
                with preview_cols[i % 4]:
                    st.markdown(f"**{name}**")
                    for code, sname in STOCK_POOLS[name]:
                        st.caption(f"  {code} {sname}")

    # ==================== TAB 8: 龙头看板 ====================
    with tabs[7]:
        st.header("🏆 各板块龙头看板")
        st.markdown("自动筛选每个板块中 AI 评分最高的龙头标的，一屏纵览全市场机会。")

        scan_df = st.session_state.get('scan_df')
        if scan_df is None or scan_df.empty:
            st.info("👆 请先在「🐉 一键擒龙」中扫描一次，数据将自动汇聚至此。")
        else:
            st.caption(f"数据来源: 上次扫描 ({st.session_state.get('scan_time', '-')})")

            # 每板块取 buy_score 最高者
            leaders = scan_df.loc[scan_df.groupby('sector')['buy_score'].idxmax()].copy()
            leaders = leaders.sort_values('buy_score', ascending=False).reset_index(drop=True)

            # 统计
            buy_count = len(leaders[leaders['signal'].isin(['STRONG_BUY', 'BUY', 'WEAK_BUY'])])
            total_sectors = len(leaders)
            st.markdown(f"> 覆盖 **{total_sectors}** 个板块 | 买入信号 **{buy_count}** 个 | "
                        f"平均评分 **{leaders['buy_score'].mean():.0f}** 分")

            st.divider()

            # 卡片布局
            cols = st.columns(4)
            for i, (_, row) in enumerate(leaders.iterrows()):
                pct = getProgressPct(row.get('buy_score', 0), 80)
                sig = row.get('signal', '')
                sig_emoji = {'STRONG_BUY': '🟢', 'BUY': '🔵', 'WEAK_BUY': '⚪',
                             'NEUTRAL': '⬜', 'WEAK_SELL': '🟡', 'SELL': '🟠',
                             'STRONG_SELL': '🔴'}.get(sig, '⬜')

                col_idx = i % 4
                with cols[col_idx]:
                    with st.container(border=True):
                        st.markdown(f"**{sig_emoji} {row.get('name', '-')}**")
                        st.caption(f"📌 {row.get('sector', '-')}")
                        st.progress(pct / 100, text=f"{row.get('buy_score', 0)}分")
                        st.caption(f"💰 {row.get('price', '-')} | "
                                   f"{row.get('change%', 0):+.1f}% | "
                                   f"RSI {row.get('rsi', '-')}")
                        strat = row.get('buy_strategy', '')
                        if strat:
                            st.caption(f"🟢 {strat}")
                        sell_s = row.get('sell_strategy', '')
                        if sell_s:
                            st.caption(f"🔴 {sell_s}")

            st.divider()

            # 一览表
            with st.expander("📊 龙头上榜清单（完整表格）"):
                leader_cols = ['sector', 'name', 'code', 'price', 'change%',
                               'signal', 'buy_score', 'rsi', 'ret_20d',
                               'buy_strategy', 'sell_strategy']
                avail_lc = [c for c in leader_cols if c in leaders.columns]
                st.dataframe(
                    leaders[avail_lc].style.background_gradient(
                        subset=['buy_score'], cmap='RdYlGn'),
                    use_container_width=True, height=400)

    # ==================== TAB 9: 资金趋势 ====================
    with tabs[8]:
        st.header("💰 资金趋势动态监控")
        st.markdown("跟踪主力资金持续流入且走势向上的标的，发现资金驱动的趋势机会。")

        scan_df2 = st.session_state.get('scan_df')
        if scan_df2 is None or scan_df2.empty:
            st.info("👆 请先在「🐉 一键擒龙」中扫描一次，数据将自动汇聚至此。")
        else:
            st.caption(f"数据来源: 上次扫描 ({st.session_state.get('scan_time', '-')})")

            # 筛选：资金流入 + 趋势向上
            fund_trend = scan_df2.copy()

            # 资金面评分
            if 'fund_flow' in fund_trend.columns:
                fund_trend['资金趋势'] = fund_trend['fund_flow'].apply(
                    lambda x: '🟢 大幅流入' if x > 3 else '🔵 流入' if x > 0 else
                    '⚪ 平衡' if x == 0 else '🟡 流出' if x > -3 else '🔴 大幅流出')
            else:
                fund_trend['资金趋势'] = '暂无数据'

            # 趋势面评分（结合20日涨幅和均线状态）
            def trend_label(row):
                ret = row.get('ret_20d', 0) or 0
                ma = str(row.get('ma_status', ''))
                score = row.get('buy_score', 0)
                if ret > 10 and '多头' in ma:
                    return '🟢 强势向上'
                elif ret > 3 or ('多头' in ma and score > 10):
                    return '🔵 稳步向上'
                elif ret > 0:
                    return '⚪ 缓慢向上'
                elif ret > -5:
                    return '🟡 横盘整理'
                else:
                    return '🔴 趋势向下'

            fund_trend['趋势状态'] = fund_trend.apply(trend_label, axis=1)

            # 综合评分（资金权重40% + 趋势权重60%）
            def composite_score(row):
                flow = float(row.get('fund_flow', 0) or 0)
                ret20 = float(row.get('ret_20d', 0) or 0)
                buy = float(row.get('buy_score', 0) or 0)
                flow_norm = max(-5, min(flow, 15))
                ret_norm = max(-20, min(ret20, 30))
                return round(flow_norm * 2 + ret_norm * 1.5 + buy * 0.5, 1)

            fund_trend['综合评分'] = fund_trend.apply(composite_score, axis=1)

            # 排序
            sort_mode = st.radio("排序依据", ["综合评分(资金+趋势)", "资金流入最多", "走势最强(20日涨幅)"],
                                 horizontal=True, key="fund_sort")
            if "资金" in sort_mode:
                fund_trend = fund_trend.sort_values('fund_flow', ascending=False)
            elif "走势" in sort_mode:
                fund_trend = fund_trend.sort_values('ret_20d', ascending=False)
            else:
                fund_trend = fund_trend.sort_values('综合评分', ascending=False)

            # 统计卡片
            mc1, mc2, mc3, mc4 = st.columns(4)
            inflow_count = len(fund_trend[fund_trend['fund_flow'] > 0]) if 'fund_flow' in fund_trend.columns else 0
            up_count = len(fund_trend[fund_trend['ret_20d'] > 0]) if 'ret_20d' in fund_trend.columns else 0
            gold_count = len(fund_trend[(fund_trend.get('fund_flow', 0) > 0) &
                                        (fund_trend.get('ret_20d', 0) > 0)]) if 'fund_flow' in fund_trend.columns else 0
            composite_top = fund_trend['综合评分'].max() if not fund_trend.empty else 0
            mc1.metric("🟢 资金流入", f"{inflow_count}只")
            mc2.metric("📈 趋势向上", f"{up_count}只")
            mc3.metric("⭐ 资金+趋势共振", f"{gold_count}只")
            mc4.metric("🏆 最高综合评分", f"{composite_top:.0f}")

            st.divider()

            # 筛选控件
            col_f1, col_f2 = st.columns(2)
            with col_f1:
                min_composite = st.slider("最低综合评分", -50.0, 100.0, 0.0, 5.0, key="min_comp")
            with col_f2:
                show_fund_up = st.checkbox("仅显示资金净流入", value=False, key="fund_up_only")

            if show_fund_up and 'fund_flow' in fund_trend.columns:
                fund_trend = fund_trend[fund_trend['fund_flow'] > 0]
            fund_trend = fund_trend[fund_trend['综合评分'] >= min_composite]

            # 显示表格
            fund_cols = ['code', 'name', 'sector', 'price', 'change%',
                         'signal', 'buy_score', 'ret_20d', 'ret_60d',
                         '资金趋势', '趋势状态',
                         'buy_strategy', 'sell_strategy']

            avail_fc = [c for c in fund_cols if c in fund_trend.columns]

            def color_flow(val):
                if '流入' in str(val) and '大' in str(val): return 'background-color: #c8e6c9'
                if '流入' in str(val): return 'background-color: #e8f5e9'
                if '流出' in str(val) and '大' in str(val): return 'background-color: #ffcdd2'
                if '流出' in str(val): return 'background-color: #ffebee'
                return ''

            def color_trend(val):
                if '强势' in str(val): return 'background-color: #c8e6c9; font-weight: bold'
                if '稳步' in str(val): return 'background-color: #e8f5e9'
                if '向下' in str(val): return 'background-color: #ffebee'
                return ''

            if not fund_trend.empty:
                styled_f = fund_trend[avail_fc].style.applymap(
                    color_flow, subset=['资金趋势'] if '资金趋势' in avail_fc else []
                ).applymap(
                    color_trend, subset=['趋势状态'] if '趋势状态' in avail_fc else []
                ).format({
                    'price': '{:.2f}', 'change%': '{:+.1f}%',
                    'ret_20d': '{:+.1f}%', 'ret_60d': '{:+.1f}%',
                })
                st.dataframe(styled_f, use_container_width=True, height=500)

                csv_fund = fund_trend[avail_fc].to_csv(index=False).encode('utf-8-sig')
                st.download_button("⬇️ 导出资金趋势CSV", csv_fund,
                                   f"资金趋势_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                                   "text/csv", key="dl_fund")
            else:
                st.info("当前无符合筛选条件的股票")

    # ==================== 底部 ====================
    st.divider()
    col_f1, col_f2, col_f3 = st.columns(3)
    with col_f1:
        st.caption("⚠️ 本工具仅供学习参考，不构成投资建议。股市有风险，投资需谨慎。")
    with col_f2:
        st.caption(f"数据更新: {datetime.now().strftime('%Y-%m-%d %H:%M')} | 数据源: mootdx/eastmoney/yfinance")
    with col_f3:
        st.caption("© 2026 A股智选助手 v2.0 | AI驱动量化选股")


# ==================== Mock数据生成函数 ====================
def _generate_mock_kline(code, days, stock_info):
    np.random.seed(int(hashlib.md5(code.encode()).hexdigest()[:8], 16) % 10000)
    start_price = np.random.uniform(30, 200)
    dates = pd.date_range(end=datetime.now(), periods=max(days + 50, 150), freq='B')
    trend = np.cumsum(np.random.normal(0.0003, 0.018, len(dates)))
    noise = np.random.normal(0, 0.008, len(dates)).cumsum()

    # 增加一些真实感：加入周期性波动
    cycle = np.sin(np.linspace(0, np.random.uniform(2, 6) * np.pi, len(dates))) * np.random.uniform(0.03, 0.08)
    prices = start_price * np.exp(trend + noise * 0.3 + cycle)

    data = []
    for i, date in enumerate(dates):
        base = prices[i]
        volatility = np.random.uniform(0.01, 0.03)
        high = base * (1 + abs(np.random.normal(0, volatility)))
        low = base * (1 - abs(np.random.normal(0, volatility)))
        open_p = base * (1 + np.random.normal(0, volatility * 0.5))
        close = base
        volume = int(np.random.uniform(500000, 8000000) * (1 + abs(np.random.normal(0, 0.3))))

        data.append({
            'date': date,
            'open': round(open_p, 2),
            'high': round(high, 2),
            'low': round(low, 2),
            'close': round(close, 2),
            'volume': volume,
        })

    df = pd.DataFrame(data).tail(days)
    df.set_index('date', inplace=True)
    return df


def _mock_fund_flow():
    return {
        '主力净流入(亿)': round(np.random.uniform(-5, 15), 2),
        '大单净流入(亿)': round(np.random.uniform(-3, 8), 2),
        '北向资金(亿)': round(np.random.uniform(-2, 5), 2),
        '趋势': np.random.choice(['净流入', '净流出', '平衡']),
    }


def _mock_hotspots():
    return pd.DataFrame({
        '板块名称': ['长鑫存储概念', 'MLCC概念', '先进封装(Chiplet)', '玻璃基板',
                    '氮化镓GaN', '碳化硅SiC', '磷化铟InP', '人造钻石',
                    '半导体设备', '存储芯片', 'AI算力',
                    '新能源电池', '银行', '消费电子', '光伏', '医药'],
        '今日涨幅%': [5.1, 3.5, 4.8, 2.6, 3.2, 4.1, 2.1, 1.8, 4.2, 3.8, 2.9,
                     1.5, -0.8, 1.2, 0.5, -1.5],
        '5日涨幅%': [10.2, 7.8, 9.5, 5.2, 6.8, 8.3, 4.5, 3.2, 8.5, 9.1, 6.2,
                     2.1, -1.2, 3.5, 1.8, -3.1],
        '20日涨幅%': [22.5, 16.3, 19.1, 10.5, 13.2, 17.8, 8.9, 6.5, 15.2, 18.6, 10.5,
                      4.3, -3.1, 6.2, 3.5, -5.2],
        '主力净流入(亿)': [85.3, 32.1, 55.6, 18.2, 28.5, 42.3, 12.8, 8.5, 62.5, 48.3, 35.1,
                         12.4, -10.2, 8.5, 5.2, -15.3],
        '低位信号': ['强(低位)', '中', '中', '中', '中', '强(低位)', '弱', '弱',
                     '中', '强(超跌)', '中', '弱', '中', '弱', '中', '强(超跌)'],
        '轮动预测': ['即将起飞', '持续关注', '即将起飞', '持续关注', '持续关注',
                     '即将起飞', '观望', '观望', '持续关注', '即将起飞', '持续关注',
                     '观望', '防御', '观望', '观望', '即将起飞'],
    })


def _mock_us_link():
    return {
        'NVDA': {'ticker': 'NVDA', 'yesterday_change': 5.2, 'a_stock_related': 'AI算力/GPU', 'sector': 'AI算力', 'impact': '利好A股'},
        'MU': {'ticker': 'MU', 'yesterday_change': 4.8, 'a_stock_related': '存储芯片', 'sector': '存储', 'impact': '利好A股'},
        'AMD': {'ticker': 'AMD', 'yesterday_change': 3.1, 'a_stock_related': 'AI算力/半导体', 'sector': '半导体', 'impact': '利好A股'},
        'TSM': {'ticker': 'TSM', 'yesterday_change': 2.5, 'a_stock_related': '半导体制造', 'sector': '半导体', 'impact': '利好A股'},
    }


def _mock_market():
    return {
        '上证指数': {'price': 3256.35, 'change_pct': 0.42},
        '深证成指': {'price': 10821.42, 'change_pct': 0.55},
        '创业板指': {'price': 2193.18, 'change_pct': 0.81},
    }


def _mock_limit_up():
    return {'limit_up': 87, 'limit_down': 12}


def _mock_financial():
    return {'最新价': 138.5, 'PE(TTM)': 32.5, 'PB': 5.2, '总市值(亿)': 580, '52周最高': 165.0, '52周最低': 85.2}


if __name__ == "__main__":
    main()
